"""Convert training markdown files to .docx format with proper formatting."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TRAINING_DIR = Path(__file__).parent
DOCS_DIR = TRAINING_DIR.parent / "docs"

MD_FILES = [
    "01-field-technician-guide.md",
    "02-dispatcher-guide.md",
    "03-admin-guide.md",
    "04-onboarding-checklist.md",
]

DOCS_MD_FILES = [
    "06-geofencing-location-stack.md",
]

BRAND_NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
BRAND_TEAL  = RGBColor(0x08, 0x91, 0xB2)
BRAND_DARK  = RGBColor(0x1F, 0x29, 0x37)

def style_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_DARK
    doc.styles['Normal'].paragraph_format.space_after = Pt(6)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = BRAND_NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = BRAND_TEAL
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = BRAND_NAVY
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND_DARK
        p.paragraph_format.space_before = Pt(8)
    return p

def add_table_from_md(doc, lines):
    """Parse a markdown table block into a Word table."""
    rows = [l for l in lines if l.strip() and not re.match(r'^\|[-| :]+\|$', l.strip())]
    if not rows:
        return
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(parsed):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text
            run = cell.paragraphs[0].runs
            if run:
                run[0].font.size = Pt(10)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = BRAND_NAVY
    doc.add_paragraph()

def inline_format(paragraph, text):
    """Add a run with bold/inline-code markup stripped but bold applied."""
    # Split on **bold** markers
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        # strip backtick code spans
        part_clean = re.sub(r'`([^`]+)`', r'\1', part)
        run = paragraph.add_run(part_clean)
        run.font.size = Pt(11)
        if idx % 2 == 1:  # was inside ** **
            run.bold = True

def convert_md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    style_doc(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4); i += 1; continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3); i += 1; continue
        if line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2); i += 1; continue
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1); i += 1; continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            doc.add_paragraph('─' * 60)
            i += 1; continue

        # Table block
        if line.strip().startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # Code block
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            continue

        # Bullet list
        if re.match(r'^[-*] ', line):
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:].strip()
            inline_format(p, content)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1; continue

        # Checkbox list
        if re.match(r'^- \[[ x]\] ', line):
            p = doc.add_paragraph(style='List Bullet')
            checked = 'x' in line[3:5]
            content = ('☑ ' if checked else '☐ ') + line[6:].strip()
            inline_format(p, content)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1; continue

        # Numbered list
        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            content = re.sub(r'^\d+\. ', '', line).strip()
            inline_format(p, content)
            i += 1; continue

        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1; continue

        # Blank line
        if not line.strip():
            i += 1; continue

        # Normal paragraph
        p = doc.add_paragraph()
        inline_format(p, line.strip())
        i += 1

    doc.save(str(docx_path))
    print(f"  ✓  {docx_path.name}")

def main():
    print("Converting training guides to .docx ...\n")
    for filename in MD_FILES:
        md_path = TRAINING_DIR / filename
        if not md_path.exists():
            print(f"  SKIP  {filename} (not found)")
            continue
        docx_name = md_path.stem + ".docx"
        docx_path = TRAINING_DIR / docx_name
        convert_md_to_docx(md_path, docx_path)

    print("\nConverting docs to .docx ...\n")
    for filename in DOCS_MD_FILES:
        md_path = DOCS_DIR / filename
        if not md_path.exists():
            print(f"  SKIP  {filename} (not found)")
            continue
        docx_name = md_path.stem + ".docx"
        docx_path = DOCS_DIR / docx_name
        convert_md_to_docx(md_path, docx_path)

    print("\nDone.")

if __name__ == "__main__":
    main()
